from django.shortcuts import render, get_object_or_404, redirect
import requests
from .models import BookContent
# import redirect
from django.shortcuts import redirect, HttpResponseRedirect
from django.core.paginator import Paginator
from .forms import DocumentForm
from .models import Document
# from docx import Document
from docx import Document as DocxDocument
from django.views.decorators.csrf import csrf_protect
import os
from django.http import HttpResponse

# def home(request):
#     content = BookContent.objects.all()
#     # get title from BookContent
#     id = BookContent.objects.values_list('id', flat=True)
#     title = BookContent.objects.values_list('title', flat=True)
#     content = BookContent.objects.values_list('content', flat=True)
    

#     # print(title[0])
#     # print(content[0])

#     # save to dbase
#     # r = BookContent(title=title, content=content)
#     # r.save()

#     context ={
#         'id': id,
#         'titles': title,
#         'contents': content,
#         'is_dark_theme': request.session.get('is_dark_theme', False),
#     }
#     return render(request, 'home2.html', context)


def change_theme(request, **kwargs):
    if 'is_dark_theme' in request.session:
        request.session['is_dark_theme'] = not request.session['is_dark_theme']
    else:
        request.session['is_dark_theme'] = True
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))





def process_document(file_path):
    doc = DocxDocument(file_path)
    document_contents = ""
    for paragraph in doc.paragraphs:
        document_contents += paragraph.text + "\n"
    return document_contents

def upload_document(request):
   
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = Document(file=request.FILES['file'])
            document.save()
            # book_content = BookContent(title=title, content=content)
            # book_content.save()
            # Process the document's contents
            document_contents = process_document(document.file.path)
            # Pass both document_contents and book_content_details to the template
            return render(request, 'home.html')
        
    # Include book_content_details in the context for the GET request scenario as well
    return render(request, 'home.html')


# get all the documents details stored in the database
def get_all_contents_and_documents(request):
    # Existing code to query BookContent and Document instances
    book_contents = BookContent.objects.all()
    book_content_details = [{
        'title': book_content.title,
        'content': book_content.content,
    } for book_content in book_contents]

    # print('Book contents: ', book_contents)
    documents = Document.objects.all()
    document_details = [{
        'file_name': os.path.splitext(os.path.basename(document.file.name))[0],
        'file_url': document.file.url,
        'pk': document.pk,
    } for document in documents]

    # Combine both lists into a single context
    context = {
        'book_contents': book_content_details,
        'documents': document_details,
    }

    print(context)

    # Return the render call with the context
    return render(request, 'home.html', context)


def book_detail(request, book_id):
    book = get_object_or_404(Document, pk=book_id)
    print("BOOK",book.pk)
    book_contents = BookContent.objects.all()
    book_content_details = [{
        'title': book_content.title,
        'content': book_content.content,
    } for book_content in book_contents]
    context = {
        'book_contents': book_content_details,
        'book': book
    }
    print(context)
    return render(request, 'bookDetail.html', context)

# def document_detail(request, id):
#     document = get_object_or_404(BookContent, pk=id)
#     # Assuming the document file is a text file for simplicity
#     try:
#         with document.file.open('r') as file:
#             content = file.read()
#     except Exception as e:
#         content = f"Error reading file: {str(e)}"
    
#     return render(request, 'bookDetail.html', {'document': document, 'content': content})

# def serve_file(request, file_name):
#     document = get_object_or_404(Document, filename=file_name)
#     with open(document.file.path, 'rb') as file:
#         response = HttpResponse(file.read(), content_type="application/octet-stream")
#         response['Content-Disposition'] = 'inline; filename=' + document.filename
#         return response


